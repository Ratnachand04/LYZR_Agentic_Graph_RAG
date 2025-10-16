#!/usr/bin/env python3
"""
Simplified document processor without problematic dependencies.
"""

import logging
import mimetypes
import re
from pathlib import Path
from typing import Dict, Optional, Union, List
import io

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Simplified document processor supporting basic formats.
    """
    
    def __init__(self):
        """Initialize document processor."""
        # Supported file extensions
        self.supported_extensions = {
            # Text formats
            '.txt', '.md', '.markdown', '.rst', '.csv',
            
            # Web/Code formats
            '.html', '.htm', '.xml', '.json', '.yaml', '.yml',
            '.py', '.js', '.java', '.cpp', '.c', '.h', '.css', '.sql',
            
            # Document formats (basic support)
            '.pdf', '.docx', '.doc', '.rtf',
            
            # Image formats (for OCR)
            '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp',
            
            # Other formats
            '.log', '.cfg', '.ini', '.conf'
        }
        
        # Check for optional dependencies
        self.has_pdf_support = self._check_pdf_support()
        self.has_docx_support = self._check_docx_support()
        self.has_ocr_support = self._check_ocr_support()
    
    def _check_pdf_support(self) -> bool:
        """Check if PDF processing is available."""
        try:
            import PyPDF2
            return True
        except ImportError:
            try:
                import fitz  # PyMuPDF
                return True
            except ImportError:
                logger.warning("PDF support not available. Install PyPDF2 or PyMuPDF for PDF processing.")
                return False
    
    def _check_docx_support(self) -> bool:
        """Check if DOCX processing is available."""
        try:
            import docx
            return True
        except ImportError:
            logger.warning("DOCX support not available. Install python-docx for DOCX processing.")
            return False
    
    def _check_ocr_support(self) -> bool:
        """Check if OCR processing is available."""
        try:
            import pytesseract
            from PIL import Image
            
            # Test if tesseract is actually accessible
            try:
                pytesseract.get_tesseract_version()
                return True
            except (pytesseract.TesseractNotFoundError, FileNotFoundError):
                # Try common Windows installation paths
                import os
                import platform
                
                if platform.system() == "Windows":
                    common_paths = [
                        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                        r"C:\Tesseract-OCR\tesseract.exe"
                    ]
                    
                    for path in common_paths:
                        if os.path.exists(path):
                            logger.info(f"Found Tesseract at {path}")
                            pytesseract.pytesseract.tesseract_cmd = path
                            try:
                                pytesseract.get_tesseract_version()
                                logger.info("Tesseract OCR configured successfully")
                                return True
                            except Exception as e:
                                logger.warning(f"Tesseract at {path} failed test: {e}")
                                continue
                
                logger.warning("Tesseract OCR engine not found. Please install Tesseract OCR.")
                logger.info("Installation instructions:")
                logger.info("Windows: winget install --id UB-Mannheim.TesseractOCR")
                logger.info("Or download from: https://github.com/UB-Mannheim/tesseract/wiki")
                logger.info("After installation, restart your terminal or add to PATH")
                return False
            except Exception as e:
                logger.warning(f"OCR test failed: {e}")
                return False
                
        except ImportError:
            logger.warning("OCR support not available. Install: pip install pytesseract pillow")
            return False
    
    def is_supported(self, file_path: Union[str, Path]) -> bool:
        """Check if file format is supported."""
        path = Path(file_path)
        return path.suffix.lower() in self.supported_extensions
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return sorted(list(self.supported_extensions))
    
    async def extract_text(self, 
                          file_path: Union[str, Path], 
                          encoding: str = 'utf-8') -> Dict[str, Union[str, Dict]]:
        """
        Extract text from supported document formats.
        
        Returns:
            Dict containing text, metadata, method, success status
        """
        path = Path(file_path)
        
        if not path.exists():
            return {
                'text': '',
                'metadata': {},
                'method': 'none',
                'success': False,
                'error': f'File not found: {path}'
            }
        
        if not self.is_supported(path):
            return {
                'text': '',
                'metadata': {},
                'method': 'none', 
                'success': False,
                'error': f'Unsupported file format: {path.suffix}'
            }
        
        # Get file metadata
        metadata = {
            'file_path': str(path.absolute()),
            'file_size': path.stat().st_size,
            'file_extension': path.suffix.lower(),
            'mime_type': mimetypes.guess_type(str(path))[0]
        }
        
        try:
            # Extract text based on file type
            extension = path.suffix.lower()
            
            if extension in ['.html', '.htm', '.xml']:
                result = await self._extract_html(path, encoding)
            elif extension == '.json':
                result = await self._extract_json(path, encoding)
            elif extension in ['.pdf'] and self.has_pdf_support:
                result = await self._extract_pdf(path)
            elif extension in ['.docx', '.doc'] and self.has_docx_support:
                result = await self._extract_docx(path)
            elif extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp']:
                if self.has_ocr_support:
                    result = await self._extract_image_ocr(path)
                else:
                    result = {
                        'text': '',
                        'method': 'image_no_ocr',
                        'success': False,
                        'error': 'OCR not available. Please install Tesseract OCR to extract text from images.',
                        'suggestion': 'Run: winget install --id UB-Mannheim.TesseractOCR (Windows)'
                    }
            else:
                # Plain text files (fallback)
                result = await self._extract_text_file(path, encoding)
            
            # Add metadata to result
            result['metadata'] = {**metadata, **result.get('metadata', {})}
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from {path}: {e}")
            return {
                'text': '',
                'metadata': metadata,
                'method': 'error',
                'success': False,
                'error': str(e)
            }
    
    async def _extract_text_file(self, path: Path, encoding: str) -> Dict:
        """Extract text from plain text files."""
        try:
            with open(path, 'r', encoding=encoding) as f:
                text = f.read()
            
            return {
                'text': text,
                'method': 'text_file',
                'success': True,
                'metadata': {'encoding': encoding}
            }
        except UnicodeDecodeError:
            # Try common encodings
            for enc in ['utf-8', 'latin1', 'cp1252', 'ascii']:
                try:
                    with open(path, 'r', encoding=enc) as f:
                        text = f.read()
                    return {
                        'text': text,
                        'method': 'text_file',
                        'success': True,
                        'metadata': {'encoding': enc}
                    }
                except UnicodeDecodeError:
                    continue
            
            raise ValueError(f"Unable to decode text file with common encodings")
    
    async def _extract_html(self, path: Path, encoding: str) -> Dict:
        """Extract text from HTML/XML files."""
        try:
            with open(path, 'r', encoding=encoding) as f:
                content = f.read()
            
            # Simple HTML text extraction without BeautifulSoup
            import re
            
            # Remove script and style content
            content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
            content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL | re.IGNORECASE)
            
            # Remove HTML tags
            text = re.sub(r'<[^>]+>', ' ', content)
            
            # Clean up whitespace
            text = re.sub(r'\s+', ' ', text).strip()
            
            return {
                'text': text,
                'method': 'simple_html',
                'success': True,
                'metadata': {'encoding': encoding}
            }
        
        except Exception as e:
            logger.warning(f"HTML extraction failed for {path}: {e}")
            return await self._extract_text_file(path, encoding)
    
    async def _extract_json(self, path: Path, encoding: str) -> Dict:
        """Extract text from JSON files."""
        try:
            import json
            with open(path, 'r', encoding=encoding) as f:
                data = json.load(f)
            
            # Convert JSON to readable text
            def json_to_text(obj, depth=0):
                if depth > 5:  # Prevent infinite recursion
                    return str(obj)
                
                text_parts = []
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, (dict, list)):
                            text_parts.append(f"{key}: {json_to_text(value, depth+1)}")
                        else:
                            text_parts.append(f"{key}: {str(value)}")
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        text_parts.append(json_to_text(item, depth+1))
                else:
                    text_parts.append(str(obj))
                return " ".join(text_parts)
            
            text = json_to_text(data)
            
            return {
                'text': text,
                'method': 'json_parser',
                'success': True,
                'metadata': {'encoding': encoding, 'json_keys': len(data) if isinstance(data, dict) else 0}
            }
        
        except Exception as e:
            logger.warning(f"JSON extraction failed for {path}: {e}")
            return await self._extract_text_file(path, encoding)
    
    async def _extract_pdf(self, path: Path) -> Dict:
        """Extract text from PDF files."""
        try:
            text = ""
            method = "unknown"
            
            # Try PyMuPDF first (better performance)
            try:
                import fitz
                doc = fitz.open(str(path))
                text_parts = []
                for page in doc:
                    text_parts.append(page.get_text())
                text = "\n".join(text_parts)
                doc.close()
                method = "pymupdf"
            except ImportError:
                # Fall back to PyPDF2
                try:
                    import PyPDF2
                    with open(path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        text_parts = []
                        for page in reader.pages:
                            text_parts.append(page.extract_text())
                        text = "\n".join(text_parts)
                    method = "pypdf2"
                except ImportError:
                    raise Exception("No PDF processing library available")
            
            # Clean up text
            text = re.sub(r'\s+', ' ', text).strip()
            
            if not text:
                return {
                    'text': '',
                    'method': method,
                    'success': False,
                    'error': 'No extractable text found in PDF'
                }
            
            return {
                'text': text,
                'method': method,
                'success': True,
                'metadata': {'pdf_pages': text.count('\n\n') + 1}
            }
        
        except Exception as e:
            logger.warning(f"PDF extraction failed for {path}: {e}")
            return {
                'text': '',
                'method': 'pdf_error',
                'success': False,
                'error': str(e)
            }
    
    async def _extract_docx(self, path: Path) -> Dict:
        """Extract text from DOCX files."""
        try:
            import docx
            
            doc = docx.Document(str(path))
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            
            if not text:
                return {
                    'text': '',
                    'method': 'python_docx',
                    'success': False,
                    'error': 'No extractable text found in DOCX'
                }
            
            return {
                'text': text,
                'method': 'python_docx',
                'success': True,
                'metadata': {
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables)
                }
            }
        
        except Exception as e:
            logger.warning(f"DOCX extraction failed for {path}: {e}")
            return {
                'text': '',
                'method': 'docx_error',
                'success': False,
                'error': str(e)
            }
    
    async def _extract_image_ocr(self, path: Path) -> Dict:
        """Extract text from images using OCR."""
        try:
            import pytesseract
            from PIL import Image
            
            # Open and process image
            with Image.open(str(path)) as img:
                # Convert to RGB if necessary
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Extract text using OCR
                text = pytesseract.image_to_string(img)
            
            # Clean up text
            text = re.sub(r'\s+', ' ', text).strip()
            
            if not text:
                return {
                    'text': '',
                    'method': 'tesseract_ocr',
                    'success': False,
                    'error': 'No text detected in image'
                }
            
            return {
                'text': text,
                'method': 'tesseract_ocr',
                'success': True,
                'metadata': {
                    'image_size': f"{img.width}x{img.height}" if 'img' in locals() else 'unknown',
                    'ocr_confidence': 'unknown'  # Could be enhanced with confidence scores
                }
            }
        
        except Exception as e:
            logger.warning(f"OCR extraction failed for {path}: {e}")
            return {
                'text': '',
                'method': 'ocr_error', 
                'success': False,
                'error': str(e)
            }


# Convenience function for simple usage
async def extract_text_from_file(file_path: Union[str, Path], **kwargs) -> str:
    """
    Simple function to extract text from a file.
    """
    processor = DocumentProcessor()
    result = await processor.extract_text(file_path)
    
    if result['success']:
        return result['text']
    else:
        logger.error(f"Failed to extract text: {result.get('error', 'Unknown error')}")
        return ""